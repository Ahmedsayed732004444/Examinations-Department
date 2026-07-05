import docx

path = r"C:\Users\AIO\Downloads\مقاييس التميز الشخصي والاحتراف المهني.docx"
doc = docx.Document(path)

with open('docx_raw.txt', 'w', encoding='utf-8') as f:
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ''
            f.write(f"[{i:04d}][{style}] {text}\n")

print(f"Total paragraphs: {len(doc.paragraphs)}")
